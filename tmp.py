import os
import uuid
import random
from pathlib import Path

# --- 必要的第三方库 ---
from PIL import Image, ImageDraw  # 用于创建图片
import numpy as np  # 用于生成图片噪点
import docx  # 用于创建Word文档
from openpyxl import Workbook  # 用于创建Excel文件


def create_random_image(filepath, width, height):
    """创建一个带有随机噪点的彩色图片。"""
    try:
        # 创建一个三通道 (RGB) 的随机像素数组
        array = np.random.randint(0, 256, (height, width, 3), dtype=np.uint8)
        image = Image.fromarray(array, 'RGB')

        # 可以在图片上画一些随机线条，让它看起来不那么单调
        draw = ImageDraw.Draw(image)
        for _ in range(10):
            x1, y1 = random.randint(0, width), random.randint(0, height)
            x2, y2 = random.randint(0, width), random.randint(0, height)
            color = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
            draw.line((x1, y1, x2, y2), fill=color, width=random.randint(1, 3))

        image.save(filepath)
    except Exception as e:
        print(f"创建图片失败 {filepath}: {e}")


def create_simple_docx(filepath, content):
    """创建一个内容简单的Word文档。"""
    try:
        document = docx.Document()
        document.add_heading(f'文档标题 - {Path(filepath).stem}', level=1)
        document.add_paragraph(content)
        document.save(filepath)
    except Exception as e:
        print(f"创建Word文档失败 {filepath}: {e}")


def create_simple_xlsx(filepath):
    """创建一个包含简单数据的Excel文件。"""
    try:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Sheet1"

        # 填充一些示例数据
        sheet["A1"] = "月份"
        sheet["B1"] = "销售额"
        for i in range(2, 10):
            sheet[f"A{i}"] = f"{i - 1}月"
            sheet[f"B{i}"] = random.randint(1000, 5000)

        workbook.save(filepath)
    except Exception as e:
        print(f"创建Excel文件失败 {filepath}: {e}")


def main():
    """主函数：创建所有演示文件。"""
    print("开始为DeepShell演示准备测试文件...")

    # 1. 定位桌面，并创建目标文件夹
    try:
        desktop_path = Path.home() / "Desktop"
        target_folder = desktop_path / "季度报告-待处理"
        target_folder.mkdir(exist_ok=True)
        print(f"成功创建或找到目标文件夹: {target_folder}")
    except Exception as e:
        print(f"错误：无法在桌面上创建目标文件夹: {e}")
        return

    # 2. 定义要生成的文件类型和数量
    num_images = 100
    num_docs = 50
    image_extensions = ['.jpg', '.png', '.gif']
    doc_extensions = ['.docx', '.xlsx']

    # 3. 生成图片文件
    print(f"\n正在生成 {num_images} 张图片...")
    for i in range(num_images):
        # 使用UUID生成随机、唯一的文件名
        filename = f"{uuid.uuid4().hex[:8]}{random.choice(image_extensions)}"
        filepath = target_folder / filename

        # 随机尺寸
        width, height = random.randint(200, 800), random.randint(200, 800)

        create_random_image(filepath, width, height)
        print(f"  ({i + 1}/{num_images}) 已创建: {filename}")
    print("图片文件生成完成。")

    # 4. 生成文档文件
    print(f"\n正在生成 {num_docs} 个文档...")
    for i in range(num_docs):
        ext = random.choice(doc_extensions)
        filename = f"report_{uuid.uuid4().hex[:6]}{ext}"
        filepath = target_folder / filename

        if ext == '.docx':
            content = f"这是第 {i + 1} 号报告的示例内容。\n生成时间: {uuid.uuid4()}"
            create_simple_docx(filepath, content)
        elif ext == '.xlsx':
            create_simple_xlsx(filepath)

        print(f"  ({i + 1}/{num_docs}) 已创建: {filename}")
    print("文档文件生成完成。")

    print(f"\n🎉 所有文件已成功生成在您的桌面上的 '{target_folder.name}' 文件夹中！")


if __name__ == "__main__":
    main()