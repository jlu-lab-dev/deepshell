import os
import platform
import subprocess
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def open_document(file_path):
    """
    自动用默认程序打开文件（跨平台支持）
    """
    system = platform.system()
    try:
        if system == "Windows":
            os.startfile(file_path)  # Windows原生方式
        elif system == "Darwin":
            subprocess.run(["open", file_path])  # macOS
        else:
            subprocess.run(["xdg-open", file_path])  # Linux
        return True
    except Exception as e:
        print(f"文件打开失败: {e}")
        return False


def open_report(report_path):
    """
    一键生成并打开报告
    """
    try:
        if os.path.exists(report_path):
            print(f"报告已生成: {report_path}")
            if open_document(report_path):
                return report_path
            else:
                # 如果自动打开失败，提示手动打开
                print(f"请手动打开文件: {report_path}")
                return report_path
        else:
            raise FileNotFoundError("报告生成失败")

    except Exception as e:
        print(f"操作失败: {e}")
        return None


# ===== 文档样式配置 =====
def set_chinese_font(doc):
    # 全局设置中英文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题字体设置
    for level in range(1, 6):
        doc.styles[f'Heading {level}'].font.name = '微软雅黑'
        doc.styles[f'Heading {level}']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def generate_analyst_report(result):
    """
    将分析结果保存为格式化的Word文档(.docx)

    参数:
        result: 包含分析结果的字符串或字典（结构化数据）

    返回:
        docx_path: 生成的临时文件路径
    """
    # 创建临时目录和文档对象
    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, f"Analysis_report{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
    doc = Document()
    set_chinese_font(doc)

    # ===== 内容写入逻辑 =====
    def add_heading(text, level=1):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_bullet(text):
        """添加项目符号"""
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(text)

    # 主标题
    add_heading('数据分析报告', level=0)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}")

    # 处理不同类型的结果输入
    if isinstance(result, str):
        # 纯文本处理
        for line in result.strip().split('\n'):
            if line.startswith(('【', '•', '▶')):
                add_heading(line.replace('【', '').replace('】', '').strip(), level=2)
            elif line.startswith(('-', '✓', '↑', '↓', '→', '!')):
                add_bullet(line)
            else:
                doc.add_paragraph(line)
    elif isinstance(result, dict):
        # 结构化数据处理
        for section, content in result.items():
            add_heading(section, level=1)
            if isinstance(content, list):
                for item in content:
                    add_bullet(str(item))
            else:
                doc.add_paragraph(str(content))

    # ===== 保存文档 =====
    doc.save(docx_path)
    return docx_path

